import json
from pathlib import Path

from docx import Document

from wxdoc_desktop.service import ConversionRequest, convert_document


def test_docx_conversion_writes_document_and_reports(tmp_path: Path):
    source = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("项目概述", level=1)
    document.add_paragraph("这是一段用于独立程序回归测试的正文。")
    document.save(source)

    output = tmp_path / "sample_WX格式.docx"
    result = convert_document(ConversionRequest(source, output))

    assert result.output_path.is_file()
    assert result.report_path.is_file()
    assert result.json_report_path.is_file()
    assert result.status in {"completed", "review"}
    report = json.loads(result.json_report_path.read_text(encoding="utf-8"))
    assert report["application"]["offline"] is True
    assert report["application"]["engine_version"] == "0.12.7"
    assert report["template_finalizer"]["style_audit"]["unexpected_styles"] == []
